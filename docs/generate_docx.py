#!/usr/bin/env python3
"""Generates ReservationSystem documentation as DOCX."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "/home/dzeyd/ReservationSystem/docs/ReservationSystem-Dokumentacja.docx"

def _shd(color):
    e = OxmlElement("w:shd")
    e.set(qn("w:fill"), color); e.set(qn("w:val"), "clear")
    return e

def _pBdr(color="6496C8"):
    b = OxmlElement("w:pBdr")
    for s in ("top", "left", "bottom", "right"):
        ch = OxmlElement(f"w:{s}")
        ch.set(qn("w:val"), "single"); ch.set(qn("w:sz"), "6")
        ch.set(qn("w:color"), color); ch.set(qn("w:space"), "4")
        b.append(ch)
    return b

def _shd_cell(cell, color):
    cell._tc.get_or_add_tcPr().append(_shd(color))

def header_row(table, texts, color="1E3C78"):
    for i, t in enumerate(texts):
        c = table.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(t); r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shd_cell(c, color)

def add_row(table, texts, shade=False):
    row = table.add_row()
    for i, t in enumerate(texts):
        c = row.cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(str(t)); r.font.size = Pt(9)
        if i == 0: r.bold = True
        if shade: _shd_cell(c, "F0F4FF")

def code(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(line)
        run.font.name = "Courier New"; run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(30, 30, 30)
        run._element.get_or_add_rPr().append(_shd("F0F0F5"))

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

# ── build ────────────────────────────────────────────

def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ────── TITLE PAGE ──────
    for _ in range(8): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ReservationSystem"); r.font.size = Pt(32)
    r.bold = True; r.font.color.rgb = RGBColor(20, 60, 120)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Asset Reservation Management System")
    r.font.size = Pt(16); r.font.color.rgb = RGBColor(60, 60, 60)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Documentation - Version 1.0")
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ASP.NET Core MVC  .NET 9  SQL Server  SAML2  Hangfire")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_page_break()

    # ────── TABLE OF CONTENTS ──────
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1.  System Overview",
        "2.  Application Architecture",
        "3.  Project Structure",
        "4.  Database Schema",
        "5.  Reservation Lifecycle",
        "6.  Feedback Mechanism",
        "7.  Roles and Permissions",
        "8.  Group Access Scoping",
        "9.  Customization Guide",
        "10. Configuration and Setup",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if item.startswith("    "):
            p.paragraph_format.left_indent = Cm(1.5)
    doc.add_page_break()

    # ════════════════════════════════════
    # 1. SYSTEM OVERVIEW
    # ════════════════════════════════════
    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph(
        "ReservationSystem is a generic ASP.NET Core MVC application for managing "
        "reservations, fault reporting, and asset management. It supports role-based "
        "access control with department-level scoping and optional SAML2 SSO."
    )

    doc.add_heading("Key Concepts", level=2)
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    header_row(t, ["Concept", "Description"])
    for i, (a, b) in enumerate([
        ("Reservations", "Users reserve assets. Approval workflow: Pending -> Accepted / Rejected / Canceled."),
        ("Feedback System", "At pickup/return, users report condition. 2-day window; expired feedback alerts managers."),
        ("Fault Reporting", "Report issues on assets. Managers mark as fixed. Minor faults flagged as drivable."),
        ("User Groups", "Departments scoping manager visibility."),
    ]): add_row(t, [a, b], i % 2)
    doc.add_paragraph()

    doc.add_heading("Technology Stack", level=2)
    for x in [".NET 9 (ASP.NET Core MVC)", "Entity Framework Core + SQL Server",
              "ASP.NET Core Identity", "Sustainsys.Saml2 (SAML2 SSO)",
              "Hangfire", "ClosedXML (Excel)", "MailKit (email)"]:
        bullet(doc, x)

    doc.add_heading("Architecture", level=2)
    code(doc,
        "+--------------------------------------------------+\n"
        "|  Web Layer (MVC) - Areas: Employee, Manager, Admin |\n"
        "+--------------------------------------------------+\n"
        "|  Service / Repository Layer (DataAccess)           |\n"
        "+--------------------------------------------------+\n"
        "|  Domain Layer (Models)                             |\n"
        "+--------------------------------------------------+\n"
        "|  Utility Layer                                     |\n"
        "+--------------------------------------------------+"
    )

    doc.add_page_break()

    # ════════════════════════════════════
    # 2. RESERVATION LIFECYCLE
    # ════════════════════════════════════
    doc.add_heading("5. Reservation Lifecycle", level=1)
    code(doc,
        "                    User creates reservation\n"
        "                              |\n"
        "                       [Pending]\n"
        "                        /      \\\n"
        "              Manager approves  Manager rejects\n"
        "                      |              |\n"
        "              [Accepted]         [Rejected]\n"
        "             /          \\\n"
        "    Pickup feedback   Cancel\n"
        "      (2-day window)    |\n"
        "            |        [Canceled]\n"
        "     [Return feedback]\n"
        "       (2-day window)\n"
        "            |\n"
        "         [Done]"
    )

    doc.add_heading("Creating a Reservation", level=2)
    bullet(doc, "User browses available assets")
    bullet(doc, "Selects asset, dates, destination, accepts terms")
    bullet(doc, "System checks for date conflicts")
    bullet(doc, "Reservation saved with Pending status")
    bullet(doc, "Notifications sent via email (if configured)")

    doc.add_heading("Approval / Rejection", level=2)
    bullet(doc, "Manager reviews pending reservations (scoped to their groups)")
    bullet(doc, "Approves or rejects. User is notified via email.")

    doc.add_heading("Feedback", level=2)
    bullet(doc, "Pickup and return feedback with mileage, fuel, cleanliness, fault reporting")
    bullet(doc, "2-day expiration window; expired feedback triggers alerts")

    doc.add_page_break()

    # ════════════════════════════════════
    # 7. ROLES
    # ════════════════════════════════════
    doc.add_heading("7. Roles and Permissions", level=1)
    doc.add_paragraph("Three cumulative roles:")

    t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    header_row(t, ["Role", "Level", "Capabilities"])
    for i, (a, b, c) in enumerate([
        ("Employee", "Basic", "Browse, reserve, provide feedback"),
        ("Manager", "Elevated", "All Employee + manage assets, reservations, faults (scoped)"),
        ("Admin", "Full", "All Manager + user management, groups, full access"),
    ]): add_row(t, [a, b], i % 2)
    doc.add_paragraph()

    doc.add_heading("Customization", level=2)
    doc.add_paragraph(
        "This project is designed as a generic reservation system. The Asset entity "
        "uses a flexible AssetType enum. To customize for your domain:\n"
        "- Add asset types in Models/Asset.cs\n"
        "- Rename AssetTag to your identifier format\n"
        "- Add or remove properties\n"
        "- Localize UI labels"
    )

    # ── SAVE ──
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"DOCX saved to {OUTPUT}")

if __name__ == "__main__":
    build()